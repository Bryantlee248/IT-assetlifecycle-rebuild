from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\Codex Project\IT-assetlifecycle-rebuild\MVP-1开发提示词.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_code_block(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", 1)
    else:
        code = styles["CodeBlock"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_after = Pt(1)
    code.paragraph_format.line_spacing = 1.0


prompt = """请基于当前已通过复核的 MVP-0 平台底座，开始 MVP-1：元数据与资产核心开发。

开发前必须先阅读并遵守以下文档：

D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\README.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\01-产品需求说明书_PRD.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\02-MVP功能清单与用户故事.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\03-系统模块设计说明.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\05-数据模型设计说明.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\06-数据库表结构设计.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\07-API接口规格说明.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\08-模板体系设计.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\10-权限与安全设计.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\12-测试用例与验收标准.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\13-种子数据与演示数据.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\15-开发总提示词.md
D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\16-UI设计规范与前端交互范式.md

本阶段只开发 MVP-1：元数据与资产核心。

一、MVP-1 开发目标

在 MVP-0 平台底座之上，实现：

1. 资产类型树
2. 字段定义
3. 表单配置
4. 列表配置
5. 查询配置
6. 运行时元数据接口
7. 资产新增
8. 资产编辑
9. 资产详情
10. 资产列表
11. JSONB 动态字段
12. 热点字段物理列
13. 资产关系
14. 字段权限基础能力
15. 服务器、网络设备、安全设备、软件许可证预置元数据

二、本阶段明确不做

请不要实现以下功能：

- 生命周期状态机
- 生命周期动作 API
- 审批
- 审批模板
- 审批任务
- 软件许可证分配、回收、续费
- 导入导出
- 报表
- 通知
- API Token
- Webhook
- 附件
- 证书、域名、云资源轻量资产
- CMDB
- 自动发现
- BPMN 流程设计器

这些属于 MVP-2 或后续阶段。

三、后端开发要求

请在现有 Spring Boot 3 + Java 21 项目中继续开发，保持模块化单体结构。

建议新增模块包：

com.itam.metadata
com.itam.asset

必须新增 Flyway 迁移，建议从当前版本继续：

V3__metadata_asset_core.sql

必须实现以下表结构，字段以 docs/mvp/06-数据库表结构设计.md 为准：

1. asset_types
2. field_definitions
3. form_schemas
4. list_view_schemas
5. search_schemas
6. assets
7. asset_relations
8. locations（如当前实现需要位置基础数据）
9. asset_location_history（如当前实现需要位置历史）

assets 表必须包含以下物理列：

id
tenant_id
asset_no
asset_name
asset_kind
asset_type_id
lifecycle_status
owner_user_id
owner_org_id
location_id
cost_center_id
responsible_user_id
serial_no
brand
model
vendor
warranty_end_date
license_end_date
source_type
sync_source
metadata_version
status
attributes jsonb
created_by
updated_by
created_at
updated_at
deleted

必须实现部分唯一索引：

CREATE UNIQUE INDEX ux_assets_tenant_asset_no_active
ON assets(tenant_id, asset_no)
WHERE deleted = false;

CREATE UNIQUE INDEX ux_assets_tenant_serial_no_active
ON assets(tenant_id, serial_no)
WHERE deleted = false AND serial_no IS NOT NULL;

必须实现常用查询索引：

tenant_id + asset_type_id
tenant_id + lifecycle_status
tenant_id + location_id
tenant_id + warranty_end_date
tenant_id + license_end_date
attributes GIN

四、JSONB 与热点字段规则

动态字段默认保存到：

assets.attributes

但以下字段必须使用 assets 物理列，不允许只存 JSONB：

asset_no
asset_name
asset_kind
asset_type_id
lifecycle_status
owner_user_id
owner_org_id
responsible_user_id
location_id
cost_center_id
serial_no
brand
model
vendor
warranty_end_date
license_end_date

字段定义中如果配置：

unique_scope != none

则必须满足：

- 如果是热点字段，必须由数据库唯一索引保证。
- 如果不是热点字段，后端必须拒绝发布该唯一配置，提示该字段需要升格为热点字段或生成数据库索引方案。
- 不能只靠应用层判断唯一。

五、字段权限基础能力

MVP-1 必须实现字段权限基础能力，至少支持：

visible
editable
masked
exportable

字段权限必须作用于：

- 资产列表
- 资产详情
- 资产新增/编辑表单
- API 响应

本阶段不需要实现完整权限配置页面，但必须预留后端能力，并对默认角色生效：

- tenant_admin：全部字段可见，可编辑非系统字段
- asset_admin：资产字段可见，可编辑
- asset_operator：运维字段可编辑，财务/合同字段不可见或脱敏
- asset_user：只能看本人相关资产字段，敏感字段隐藏
- auditor：只读，敏感字段脱敏

如果完整 field_permission_rules 配置页面放到 MVP-3，本阶段可以先用默认规则服务实现，但接口响应必须已经体现字段过滤。

六、API 接口要求

所有接口必须遵守 docs/mvp/07-API接口规格说明.md：

- 统一响应信封
- 统一错误码
- traceId
- Bearer Token
- 租户上下文由后端解析
- 前端传入 tenant_id 不可信
- 分页格式统一

请实现以下接口：

元数据接口：

GET  /api/v1/metadata/asset-types/tree
POST /api/v1/metadata/asset-types
PUT  /api/v1/metadata/asset-types/{typeId}
PATCH /api/v1/metadata/asset-types/{typeId}/status

GET  /api/v1/metadata/asset-types/{typeId}/fields
POST /api/v1/metadata/asset-types/{typeId}/fields
PUT  /api/v1/metadata/fields/{fieldId}
PATCH /api/v1/metadata/fields/{fieldId}/status

GET /api/v1/metadata/asset-types/{typeId}/form-schema
PUT /api/v1/metadata/asset-types/{typeId}/form-schema

GET /api/v1/metadata/asset-types/{typeId}/list-view
PUT /api/v1/metadata/asset-types/{typeId}/list-view

GET /api/v1/metadata/asset-types/{typeId}/search-schema
PUT /api/v1/metadata/asset-types/{typeId}/search-schema

GET /api/v1/metadata/runtime/asset-types/{typeId}

资产接口：

GET  /api/v1/assets
POST /api/v1/assets
GET  /api/v1/assets/{assetId}
PUT  /api/v1/assets/{assetId}
DELETE /api/v1/assets/{assetId}

GET  /api/v1/assets/{assetId}/relations
POST /api/v1/assets/{assetId}/relations
DELETE /api/v1/assets/{assetId}/relations/{relationId}

注意：

- 资产新增/编辑接口不得允许直接修改 lifecycle_status。
- lifecycle_status 只能初始化为资产类型生命周期模板的初始状态；MVP-1 可以先使用默认初始状态 `planned` 或资产类型配置中的初始状态。
- 资产列表必须支持分页、关键词、资产类型、状态、责任人、组织、位置、到期日期范围等基础筛选。
- 资产详情和列表必须应用数据权限和字段权限。

七、前端开发要求

继续使用现有 Vue 3 + TypeScript + Vite + Element Plus + Pinia + Vue Router。

必须遵守：

D:\\Codex Project\\IT-assetlifecycle-rebuild\\docs\\mvp\\16-UI设计规范与前端交互范式.md

新增页面：

1. 资产类型管理页
2. 字段定义管理页
3. 表单配置页
4. 列表配置页
5. 查询配置页
6. 资产列表页
7. 资产详情页
8. 资产新增/编辑页
9. 资产关系管理区域

UI 要求：

- 企业级 B 端后台风格
- 左侧导航 + 顶栏 + 主工作区
- 高信息密度但不拥挤
- 使用 Element Plus 表格、表单、抽屉、标签、分页、树组件
- 所有页面必须有 loading、empty、error、no-permission 状态
- 动态字段表单必须由运行时元数据驱动
- 不允许为服务器、网络设备、安全设备、软件许可证分别写死独立页面
- 字段权限必须在 UI 上体现为隐藏、只读、脱敏
- 资产详情建议使用右侧生命周期时间线预留区域，但本阶段不实现生命周期动作

八、种子数据要求

请基于 docs/mvp/13-种子数据与演示数据.md，为 MVP-1 增加种子数据：

资产类型：

- 服务器
- 网络设备
- 安全设备
- 软件许可证

字段定义：

服务器字段至少包括：

brand
model
serial_no
memory_gb
cpu_spec
warranty_end_date

网络设备字段至少包括：

brand
model
serial_no
management_ip
warranty_end_date

安全设备字段至少包括：

brand
model
serial_no
security_zone
warranty_end_date

软件许可证字段至少包括：

vendor
license_model
total_quantity
license_end_date

演示资产：

- SRV-2026-0001 核心数据库服务器01
- NET-2026-0001 核心交换机01
- SEC-2026-0001 边界防火墙01
- LIC-2026-0001 PostgreSQL 企业支持授权

九、测试要求

必须先写测试，再实现功能。

请至少补充以下测试：

后端测试：

1. 租户管理员创建资产类型成功
2. 同一租户 asset_type.type_code 重复返回 409
3. A 租户不能访问 B 租户资产类型
4. 字段定义创建成功
5. unique_scope 配置在非热点字段上被拒绝
6. 表单/list/search schema 可保存和读取
7. 创建服务器资产成功，热点字段进入物理列，扩展字段进入 attributes
8. 相同 asset_no 重复创建返回 409
9. 相同 serial_no 重复创建返回 409
10. 软删除后相同 asset_no 可重新创建
11. 资产编辑不能直接修改 lifecycle_status
12. 资产列表分页返回统一 PageResult
13. 普通用户无法看到无权限字段
14. 审计员看到敏感字段脱敏
15. 资产关系可创建和删除
16. 跨租户访问资产详情返回 403 或 404

前端测试/构建：

cd frontend
npm run build

后端测试/构建：

cd backend
mvn test
mvn verify

十、验收标准

完成后必须对照 docs/mvp/12-测试用例与验收标准.md 的 MVP-1 验收项输出结果：

MVP-1 必须证明：

1. 租户管理员可以创建资产类型。
2. 租户管理员可以配置字段。
3. 租户管理员可以配置表单、列表、查询。
4. 运行时元数据接口返回完整配置。
5. 用户可以创建服务器资产。
6. 用户可以创建软件许可证资产。
7. 两种资产使用不同动态字段。
8. 资产关系可建立。
9. 租户隔离有效。
10. 字段权限在 API 和 UI 中生效。
11. 热点字段和唯一字段由数据库约束保证。
12. JSONB 动态字段可保存和读取。

十一、交付物要求

完成后请输出：

1. 本阶段新增/修改文件清单
2. 数据库迁移文件清单
3. 新增 API 清单
4. 前端新增页面清单
5. 种子数据说明
6. 测试用例清单
7. 后端 `mvn test` 和 `mvn verify` 结果
8. 前端 `npm run build` 结果
9. Docker 启动验证结果
10. MVP-1 验收报告
11. 已知遗留问题和是否阻塞 MVP-2

十二、重要约束

- 不要修改 MVP-0 已验证通过的认证、租户、权限、强制改密逻辑，除非是为了兼容 MVP-1 且必须有回归测试。
- 不要跳到 MVP-2。
- 不要实现生命周期动作。
- 不要实现审批。
- 不要把字段写死在前端页面中。
- 不要只用 JSONB 保存所有字段，热点字段必须落物理列。
- 不要只靠前端隐藏字段权限，后端响应必须过滤。
- 不要信任前端传入的 tenant_id。
- 所有关键写操作必须写审计日志。"""


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MVP-1 元数据与资产核心开发提示词")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor.from_string("1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("用于提交给 AI 开发平台的完整阶段开发指令")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_h(doc, "使用说明", 1)
    add_p(doc, "请将“正式提示词”一节中的内容整体复制给 AI 开发平台。该提示词用于启动 MVP-1，不应拆成多个不连续任务。", bold=True)
    add_bullet(doc, "前置条件：MVP-0 已通过复核，并已完成 Docker/API 冒烟验收。")
    add_bullet(doc, "目标范围：只开发元数据与资产核心。")
    add_bullet(doc, "禁止范围：不得进入生命周期、审批、导入导出、报表、Webhook、API Token 等后续阶段。")
    add_bullet(doc, "交付要求：必须提供测试结果、构建结果、Docker 验证和 MVP-1 验收报告。")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "项目", bold=True)
    set_cell_text(hdr[1], "说明", bold=True)
    set_cell_shading(hdr[0], "E8EEF5")
    set_cell_shading(hdr[1], "E8EEF5")
    for label, value in [
        ("适用阶段", "MVP-1 元数据与资产核心"),
        ("输入方式", "整段复制给 AI 开发平台"),
        ("核心产出", "资产类型、字段定义、动态表单、资产台账、资产关系、字段权限基础"),
        ("验收重点", "JSONB + 热点物理列、租户隔离、字段权限、唯一约束、动态 UI"),
    ]:
        row = table.add_row().cells
        set_cell_text(row[0], label, bold=True)
        set_cell_text(row[1], value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_page_break()
    add_h(doc, "正式提示词", 1)
    add_p(doc, "以下内容请整体复制给 AI 开发平台：", bold=True)
    add_code_block(doc, prompt)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
