from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"D:\Codex Project\IT-assetlifecycle-rebuild\MVP-0整改提示词.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def add_code_block(doc, text):
    for line in text.splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(9)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_p(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    return p


prompt = """请只修复 MVP-0 平台底座的审核问题，不要进入 MVP-1，不要实现资产、资产类型、动态字段、生命周期、审批、软件许可证、导入导出、报表、API Token、Webhook、附件、云资源、证书、域名等后续功能。

当前目标：让 MVP-0 达到可编译、可启动、可测试、可验收状态。

请基于当前代码完成以下整改。

一、修复后端 Java 编译错误

请检查并修复以下文件：

1. backend/src/main/java/com/itam/tenantadmin/OrganizationRepository.java
- 当前使用 Optional 但缺少：
  import java.util.Optional;

2. backend/src/main/java/com/itam/tenantadmin/RoleService.java
- 当前使用 Map.of 但缺少：
  import java.util.Map;

3. backend/src/main/java/com/itam/tenantadmin/RolePermissionService.java
- 当前使用 Collectors.toSet() 但缺少：
  import java.util.stream.Collectors;

4. backend/src/main/java/com/itam/platform/TenantService.java
- 当前存在非法写法：
  new java.util.Map.of(...)
- 请改为：
  Map.of(...)
- 并补充：
  import java.util.Map;

修复后必须确保后端可以通过 `mvn test`。

二、修正 Spring Security 匿名放行路径

当前 application.yml 配置了：

server.servlet.context-path: /api

因此 Spring Security 中 requestMatchers 不应写成 `/api/v1/...`。

请修改：

backend/src/main/java/com/itam/security/SecurityConfig.java

将匿名放行路径从：

/api/v1/auth/login
/api/v1/auth/refresh
/api/v1/health

修正为：

/v1/auth/login
/v1/auth/refresh
/v1/health

并确认以下地址实际可匿名访问：

POST http://localhost:8080/api/v1/auth/login
POST http://localhost:8080/api/v1/auth/refresh
GET  http://localhost:8080/api/v1/health

请补充对应的集成测试或 MockMvc 测试，防止后续回归。

三、后端强制执行 must_change_password

当前 must_change_password 只由前端约束，不满足 MVP-0 验收要求。

请新增后端拦截逻辑：

当 JWT principal 中 `mustChangePassword=true` 时：

允许访问：
- POST /v1/auth/change-password
- POST /v1/auth/logout
- GET /v1/auth/me

禁止访问其他所有接口，返回 HTTP 403。

响应必须仍使用统一信封：

{
  "code": 40300,
  "message": "请先修改初始密码",
  "data": null,
  "traceId": "..."
}

建议实现方式：

- 新增 `MustChangePasswordFilter`
- 放在 `JwtFilter` 之后
- 基于 `JwtUserPrincipal.isMustChangePassword()` 判断
- 不要依赖前端路由守卫作为唯一控制

必须补测试：

- must_change_password=true 的用户访问 `/v1/tenant/users` 返回 403
- must_change_password=true 的用户访问 `/v1/auth/change-password` 允许
- 改密成功后重新登录，`must_change_password=false`

四、强化平台域 / 租户域接口隔离

当前接口主要依赖权限码，但没有明确强制 userType。

请修改所有平台接口和租户接口的权限表达式。

平台接口：

- `/v1/platform/**`
- 必须要求：
  - principal.userType == PLATFORM
  - 且具备对应权限码

示例：

@PreAuthorize("principal.userType.name() == 'PLATFORM' and hasAuthority('tenant:list')")

租户接口：

- `/v1/tenant/**`
- 必须要求：
  - principal.userType == TENANT
  - 且具备对应权限码

示例：

@PreAuthorize("principal.userType.name() == 'TENANT' and hasAuthority('user:list')")

必须补测试：

- 平台管理员访问 `/v1/tenant/users` 返回 403
- 租户管理员访问 `/v1/platform/tenants` 返回 403
- 租户管理员访问本租户 `/v1/tenant/users` 正常

五、修复组织父节点跨租户校验

当前 OrganizationService 中父节点校验使用了：

organizationRepository.findById(parentId)

这会绕过 tenant_id 校验。

请修改：

- create 组织时校验 parentId：
  - 使用 `findByTenantIdAndId(currentTenantId, parentId)`

- update 组织时校验 parentId：
  - 使用 `findByTenantIdAndId(currentTenantId, parentId)`

如果 parentId 不属于当前租户，应返回 404 或 403。

必须补测试：

- A 租户创建组织时传入 B 租户组织 ID 作为 parentId，必须失败
- A 租户更新组织 parentId 为 B 租户组织 ID，必须失败

六、修复统一响应 traceId 缺失

当前 GlobalExceptionHandler 中参数校验异常手工构造 ApiResponse，但没有 traceId。

请修改：

backend/src/main/java/com/itam/common/exception/GlobalExceptionHandler.java

要求所有异常响应都必须包含非空 traceId。

建议增加：

- `ApiResponse.fail(resultCode, message, data)`
- 或统一构造方法

避免手写响应时遗漏 traceId。

必须补测试：

- 参数校验失败时响应包含 traceId
- 业务异常响应包含 traceId
- 认证异常响应包含 traceId
- 授权异常响应包含 traceId

七、修复审计日志 JSONB 映射

当前 audit_log.detail 数据库字段是 jsonb，但 Java 字段是 String，可能导致 PostgreSQL 写入类型错误。

请修复：

- backend/src/main/java/com/itam/audit/AuditLog.java
- backend/src/main/java/com/itam/audit/AuditLogService.java

优先采用方案 A：

- 使用 Hibernate 6 JSON 映射：
  - `@JdbcTypeCode(SqlTypes.JSON)`
  - 字段类型使用 `Map<String, Object>` 或 `Object`
- AuditLogService 直接保存结构化 detail

不优先采用方案 B：

- 将数据库字段改为 text

如果选择方案 B，必须同步修改 Flyway 脚本和设计说明。

必须补测试：

- 创建租户后 audit_log 中 detail 可正常写入
- 登录失败 audit_log 可正常写入
- 审计写入失败不能影响主业务
- 正常情况下审计不得失败

八、清理交付目录中的构建产物和依赖目录

当前工作目录中出现：

- frontend/node_modules
- frontend/dist

请补充或修正 `.gitignore`，至少包含：

frontend/node_modules/
frontend/dist/
backend/target/
*.log
.env

如果这些文件已经被纳入版本控制，请从版本控制中移除；如果未纳入，仅保持忽略即可。

九、重新执行验证

修复完成后必须执行并提供输出摘要。

后端：

cd backend
mvn test
mvn verify

前端：

cd frontend
npm run build

容器：

docker compose up -d --build
curl http://localhost:8080/api/v1/health

接口验收：

- 使用平台管理员登录
- 创建租户
- 使用租户管理员登录
- 验证首次强制改密
- 改密后访问租户用户列表
- 验证平台管理员不能访问租户接口
- 验证租户管理员不能访问平台接口
- 验证 refresh token 轮换
- 验证 logout 后 access token 失效
- 验证关键写操作产生 audit_log

十、输出修复报告

修复完成后请输出：

1. 修复文件清单
2. 每个问题的修复说明
3. 新增测试清单
4. 后端测试执行结果
5. 前端构建结果
6. Docker 启动结果
7. 接口验收结果
8. MVP-0 验收项重新对照结果
9. 是否仍有遗留问题

注意事项：

- 不要扩大范围。
- 不要进入 MVP-1。
- 不要新增资产、生命周期、审批等模块。
- 所有修改必须服务于 MVP-0 平台底座验收。
- 如果发现新的阻塞问题，请先修复 MVP-0 范围内的问题，并在报告中明确说明。"""


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.05)
    code.paragraph_format.space_after = Pt(1)
    code.paragraph_format.line_spacing = 1.0


def build_doc():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("MVP-0 平台底座整改提示词")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor.from_string("1F4D78")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("用于提交给 AI 开发平台的完整修复指令")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_h(doc, "使用说明", 1)
    add_p(doc, "请将“正式提示词”一节中的内容整体复制给 AI 开发平台。不要拆分发送，除非平台存在单次输入长度限制。", bold=True)
    add_bullet(doc, "目标是修复 MVP-0 审核问题，不是开发 MVP-1。")
    add_bullet(doc, "开发平台必须先修复编译、认证路径、强制改密、平台/租户边界和租户隔离问题。")
    add_bullet(doc, "修复完成后必须重新执行后端测试、前端构建、Docker 启动和接口验收。")

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "项目", bold=True)
    set_cell_text(hdr[1], "说明", bold=True)
    set_cell_shading(hdr[0], "E8EEF5")
    set_cell_shading(hdr[1], "E8EEF5")
    for label, value in [
        ("适用阶段", "MVP-0 平台底座整改"),
        ("输入方式", "整段复制给 AI 开发平台"),
        ("禁止范围", "不得进入资产、生命周期、审批等 MVP-1+ 功能"),
        ("验收重点", "可编译、可启动、可测试、可验收"),
    ]:
        row = table.add_row().cells
        set_cell_text(row[0], label, bold=True)
        set_cell_text(row[1], value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_page_break()
    add_h(doc, "正式提示词", 1)
    add_p(doc, "以下内容请整体复制给 AI 开发平台：", bold=True)

    add_code_block(doc, prompt)

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
